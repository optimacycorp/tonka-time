from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INITIAL_TARGETS = [
    (
        {
            "Customer initials - 811/private utilities responsibility: __________",
            "Customer initials - 811/private utilities responsibility: [[OS_INITIAL_811]]",
        },
        "Customer initials - 811/private utilities responsibility:",
        "[     ] ",
        "[[OS_INITIAL_811]]",
    ),
    (
        {
            "Customer initials - tutorial is not professional training: __________",
            "Customer initials - tutorial is not professional training: [[OS_INITIAL_TUTORIAL]]",
        },
        "Customer initials - tutorial is not professional training:",
        "[     ] ",
        "[[OS_INITIAL_TUTORIAL]]",
    ),
    (
        {
            "Customer initials - damage waiver limits: __________",
            "Customer initials - damage waiver limits: [[OS_INITIAL_DAMAGE_WAIVER]]",
        },
        "Customer initials - damage waiver limits:",
        "[     ] ",
        "[[OS_INITIAL_DAMAGE_WAIVER]]",
    ),
]

SIGNATURE_TARGETS = {
    "Customer signature: [OpenSign signature field]     Date signed: [OpenSign date field]",
    "Customer signature: [[OS_SIGNATURE_CUSTOMER]]     Date signed: [[OS_DATE_SIGNED]]",
}


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: update_agreement_anchor_template.py <template.docx>")

    target = Path(sys.argv[1])
    doc = Document(str(target))

    replaced_initials = 0
    replaced_signature = False

    for para in list(doc.paragraphs):
        text = para.text.strip()

        for candidates, label_text, display_text, anchor_text in INITIAL_TARGETS:
            if text not in candidates:
                continue
            para.text = label_text
            style_anchor_label(para)
            add_anchor_box_after(para, display_text, anchor_text, width_inches=1.1)
            replaced_initials += 1
            break

        if text in SIGNATURE_TARGETS:
            para.text = ""
            add_signature_row_after(para)
            replaced_signature = True

    if replaced_initials != len(INITIAL_TARGETS):
        raise SystemExit(
            f"Expected to replace {len(INITIAL_TARGETS)} initials anchors but replaced {replaced_initials}."
        )

    if not replaced_signature:
        raise SystemExit("Missing expected signature/date anchor line.")

    doc.save(str(target))
    return 0


def add_anchor_box_after(paragraph, display_text: str, anchor_text: str, width_inches: float) -> None:
    table = paragraph._parent.add_table(rows=1, cols=1, width=Inches(width_inches))
    paragraph._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders_none(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(width_inches)
    set_cell_margins(cell, 80, 40, 80, 40)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.text = ""
    visible = p.add_run(display_text)
    visible.font.name = "Courier New"
    visible.font.size = Pt(10.5)
    hidden = p.add_run(anchor_text)
    style_hidden_anchor(hidden)


def add_signature_row_after(paragraph) -> None:
    table = paragraph._parent.add_table(rows=1, cols=4, width=Inches(7.0))
    paragraph._p.addnext(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders_none(table)
    widths = [1.35, 3.1, 1.0, 1.55]
    values = [
        "Customer signature:",
        ("________________________ ", "[[OS_SIGNATURE_CUSTOMER]]"),
        "Date signed:",
        ("____________ ", "[[OS_DATE_SIGNED]]"),
    ]
    for cell, width_inches, value in zip(table.rows[0].cells, widths, values):
        cell.width = Inches(width_inches)
        set_cell_margins(cell, 60, 20, 60, 20)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.text = ""
        if isinstance(value, tuple):
            visible = p.add_run(value[0])
            visible.font.name = "Courier New"
            visible.font.size = Pt(10.5)
            hidden = p.add_run(value[1])
            style_hidden_anchor(hidden)
        else:
            run = p.add_run(value)
            run.font.size = Pt(10.5)


def style_anchor_label(paragraph) -> None:
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        run.font.size = Pt(10.5)


def style_hidden_anchor(run) -> None:
    run.font.name = "Courier New"
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)


def set_table_borders_none(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
      borders = OxmlElement("w:tblBorders")
      tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
      element = borders.find(qn(f"w:{edge}"))
      if element is None:
          element = OxmlElement(f"w:{edge}")
          borders.append(element)
      element.set(qn("w:val"), "nil")


def set_cell_margins(cell, left: int, top: int, right: int, bottom: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
      tc_mar = OxmlElement("w:tcMar")
      tc_pr.append(tc_mar)

    for edge, value in (("left", left), ("top", top), ("right", right), ("bottom", bottom)):
      element = tc_mar.find(qn(f"w:{edge}"))
      if element is None:
          element = OxmlElement(f"w:{edge}")
          tc_mar.append(element)
      element.set(qn("w:w"), str(value))
      element.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    raise SystemExit(main())
